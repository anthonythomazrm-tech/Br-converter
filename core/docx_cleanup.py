from __future__ import annotations

import re
from dataclasses import dataclass
from typing import Optional

from docx import Document
from docx.text.paragraph import Paragraph


_STRONG_END = re.compile(r"[.!?:;]\s*$")
_BULLET_START = re.compile(r"^\s*([•\-\–\*]|(\d+[\.\)])|([a-zA-Z][\.\)])|([IVXLCDM]+[\.\)]))\s+")
_ALL_CAPS_LINE = re.compile(r"^[^a-záàâãéêíóôõúç]*[A-ZÁÀÂÃÉÊÍÓÔÕÚÇ][^a-záàâãéêíóôõúç]*$")


def _norm_spaces(s: str) -> str:
    s = s.replace("\u00A0", " ")
    s = re.sub(r"[ \t]{2,}", " ", s)
    return s.strip()


def _ends_strong(text: str) -> bool:
    return bool(_STRONG_END.search(text))


def _is_blank(p: Paragraph) -> bool:
    return _norm_spaces(p.text) == ""


def _is_in_table(p: Paragraph) -> bool:
    # python-docx doesn't expose a clean flag; this is a common/robust check.
    return p._p.getparent().tag.endswith("}tc")


def _looks_like_list_item(text: str) -> bool:
    return bool(_BULLET_START.match(text))


def _looks_like_title(p: Paragraph) -> bool:
    txt = _norm_spaces(p.text)
    if not txt:
        return False

    # If style name hints heading/title, treat as title
    style_name = (p.style.name or "").lower() if p.style is not None else ""
    if "heading" in style_name or "título" in style_name or "title" in style_name:
        return True

    # Heuristic: short line, often caps, not ending with strong punctuation
    words = txt.split()
    if len(words) <= 10 and not _ends_strong(txt):
        # a lot of caps or line looks like caps-only
        if _ALL_CAPS_LINE.match(txt):
            return True

    return False


def _same_format(a: Paragraph, b: Paragraph) -> bool:
    # Compare key formatting signals to avoid merging across different blocks
    af = a.paragraph_format
    bf = b.paragraph_format

    # Alignment changes usually mean new block
    if a.alignment != b.alignment:
        return False

    # Indentation / spacing differences hint separate paragraphs
    if (af.left_indent != bf.left_indent) or (af.first_line_indent != bf.first_line_indent):
        return False
    if (af.space_before != bf.space_before) or (af.space_after != bf.space_after):
        return False

    # Compare first run style as a weak signal (not perfect but helps)
    ar = a.runs[0] if a.runs else None
    br = b.runs[0] if b.runs else None
    if ar and br:
        if (ar.bold != br.bold) or (ar.italic != br.italic) or (ar.underline != br.underline):
            # titles often bold, body not
            return False
        if (ar.font.name != br.font.name) or (ar.font.size != br.font.size):
            return False

    return True


def _merge_text(left: str, right: str) -> str:
    left = left.rstrip()
    right = right.lstrip()

    if not left:
        return right
    if not right:
        return left

    # Hyphenation fix: "computa-\ndor" => "computador"
    if left.endswith("-") and len(left) >= 2 and left[-2].isalpha() and right[:1].isalpha():
        return left[:-1] + right

    # Normal join with a space unless punctuation already ends with space behavior
    if left.endswith(("(", "[", "{", "“", '"', "‘")):
        return left + right

    return left + " " + right


def cleanup_docx_paragraphs(docx_path: str) -> None:
    """
    Post-process a DOCX to merge 'line-paragraphs' into real paragraphs.

    Safe rules:
    - Never touch paragraphs inside tables
    - Never merge list items
    - Never merge titles/headings
    - Merge only when formatting is consistent and previous line doesn't look like paragraph end
    """
    doc = Document(docx_path)
    paras = list(doc.paragraphs)

    i = 0
    while i < len(paras) - 1:
        p = paras[i]
        n = paras[i + 1]

        # Skip protected zones
        if _is_in_table(p) or _is_in_table(n):
            i += 1
            continue

        if _is_blank(p):
            i += 1
            continue

        p_txt = _norm_spaces(p.text)
        n_txt = _norm_spaces(n.text)

        if not n_txt:
            i += 1
            continue

        # Don't merge titles/lists
        if _looks_like_title(p) or _looks_like_title(n):
            i += 1
            continue
        if _looks_like_list_item(p_txt) or _looks_like_list_item(n_txt):
            i += 1
            continue

        # If previous looks like a real paragraph end, do not merge
        if _ends_strong(p_txt):
            i += 1
            continue

        # Formatting guard
        if not _same_format(p, n):
            i += 1
            continue

        # Merge: append next paragraph text into current paragraph
        merged = _merge_text(p.text, n.text)

        # Replace p runs with a single run (clean + stable)
        p.clear()
        p.add_run(_norm_spaces(merged))

        # Remove next paragraph XML
        n._element.getparent().remove(n._element)

        # Rebuild list snapshot (doc.paragraphs updates dynamically)
        paras = list(doc.paragraphs)
        # Do NOT increment i, because we may merge multiple consecutive lines

    doc.save(docx_path)