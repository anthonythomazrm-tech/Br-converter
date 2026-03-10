from __future__ import annotations

import re
from docx import Document
from docx.shared import Pt


# Linha "fake" típica: muitos _ ou muitos - (com ou sem espaços)
_FAKE_LINE_RE = re.compile(r"^\s*([_\-]\s*){8,}\s*$")

# Detecta gatilhos de campo de resposta
_ANSWER_RE = re.compile(r"^\s*resposta\s*:?\s*$", re.IGNORECASE)


def _remove_paragraph(paragraph) -> None:
    p = paragraph._element
    p.getparent().remove(p)


def _is_fake_line(text: str) -> bool:
    return bool(_FAKE_LINE_RE.match((text or "").strip()))


def count_fake_line_paragraphs(docx_path: str) -> int:
    doc = Document(docx_path)
    return sum(1 for p in doc.paragraphs if _is_fake_line(p.text))


def looks_like_form_docx(docx_path: str, threshold: int = 6) -> bool:
    return count_fake_line_paragraphs(docx_path) >= threshold


def _insert_blank_paragraphs_after(paragraph, count: int) -> None:
    """
    Insere `count` parágrafos vazios logo após `paragraph` usando XML,
    para ficar exatamente no lugar certo (não no final do doc).
    """
    # Import aqui pra evitar importar demais no topo
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    p_elm = paragraph._element
    for _ in range(count):
        new_p = OxmlElement("w:p")
        p_elm.addnext(new_p)
        p_elm = new_p  # próximo vai entrar logo depois do anterior


def fix_form_lines_in_docx(
    docx_path: str,
    default_answer_lines: int = 6,
    aggressive: bool = True,
    min_lines: int = 2,
    max_lines: int = 20,
) -> None:
    """
    Conserta formulários (modo funcional):
    - Remove parágrafos que são só ________ ou --------
    - Se encontrar "Resposta:", remove as linhas fake logo abaixo
      e insere parágrafos vazios normais (sem tabela/sem borda)
    """
    doc = Document(docx_path)

    # Normaliza fonte padrão (opcional)
    try:
        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
    except Exception:
        pass

    i = 0
    while i < len(doc.paragraphs):
        p = doc.paragraphs[i]
        txt = (p.text or "").strip()

        # Remove linhas fake soltas
        if aggressive and _is_fake_line(txt):
            _remove_paragraph(p)
            continue  # não incrementa, lista mudou

        # Se for "Resposta:", tratar bloco
        if _ANSWER_RE.match(txt):
            j = i + 1
            removed = 0

            # Remove linhas fake subsequentes (ignorando brancos)
            while j < len(doc.paragraphs):
                nxt = doc.paragraphs[j]
                t = (nxt.text or "").strip()

                if t == "":
                    j += 1
                    continue

                if _is_fake_line(t):
                    _remove_paragraph(nxt)
                    removed += 1
                    continue

                break

            if removed > 0:
                # número inteligente: recria a mesma quantidade (com limites)
                lines_to_create = max(min_lines, min(removed, max_lines))

                # insere parágrafos vazios logo após "Resposta:"
                _insert_blank_paragraphs_after(p, lines_to_create)

                # como inserimos logo abaixo, seguimos
                i += 1
                continue

            # Se não removeu nada, mas ainda quer garantir linhas mínimas
            # (opcional: mantém default)
            # _insert_blank_paragraphs_after(p, default_answer_lines)

        i += 1

    doc.save(docx_path)